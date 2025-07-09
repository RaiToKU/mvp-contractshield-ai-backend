import os
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import logging

logger = logging.getLogger(__name__)

class ExportService:
    """导出服务，生成审查报告"""
    
    def __init__(self):
        self.export_dir = "./exports"
        os.makedirs(self.export_dir, exist_ok=True)
    
    def generate_report(self, review_data: Dict[str, Any], format_type: str = "docx") -> str:
        """生成审查报告"""
        if format_type.lower() == "docx":
            return self._generate_docx_report(review_data)
        elif format_type.lower() == "pdf":
            # 先生成DOCX，然后转换为PDF（需要额外的库如python-docx2pdf）
            docx_path = self._generate_docx_report(review_data)
            return self._convert_to_pdf(docx_path)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _generate_docx_report(self, review_data: Dict[str, Any]) -> str:
        """生成DOCX格式报告"""
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading('合同风险审查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        doc.add_heading('基本信息', level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('任务ID', str(review_data['task_id'])),
            ('合同类型', review_data['contract_type']),
            ('审查角色', review_data['role']),
            ('审查状态', review_data['status']),
            ('生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = value
        
        # 风险概览
        doc.add_heading('风险概览', level=1)
        summary = review_data['summary']
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('总风险数', str(summary['total_risks'])),
            ('高风险', str(summary['high_risks'])),
            ('中风险', str(summary['medium_risks'])),
            ('低风险', str(summary['low_risks']))
        ]
        
        for i, (key, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = value
        
        # 详细风险分析
        doc.add_heading('详细风险分析', level=1)
        
        for i, risk in enumerate(review_data['risks'], 1):
            # 风险标题
            risk_heading = doc.add_heading(f'{i}. {risk["title"]}', level=2)
            
            # 风险级别
            level_para = doc.add_paragraph()
            level_para.add_run('风险级别: ').bold = True
            level_run = level_para.add_run(risk['risk_level'])
            
            # 根据风险级别设置颜色
            if risk['risk_level'] == 'HIGH':
                level_run.font.color.rgb = self._get_rgb_color(255, 0, 0)  # 红色
            elif risk['risk_level'] == 'MEDIUM':
                level_run.font.color.rgb = self._get_rgb_color(255, 165, 0)  # 橙色
            else:
                level_run.font.color.rgb = self._get_rgb_color(0, 128, 0)  # 绿色
            
            # 条款编号
            if risk['clause_id']:
                clause_para = doc.add_paragraph()
                clause_para.add_run('条款编号: ').bold = True
                clause_para.add_run(risk['clause_id'])
            
            # 风险描述
            summary_para = doc.add_paragraph()
            summary_para.add_run('风险描述: ').bold = True
            doc.add_paragraph(risk['summary'])
            
            # 应对建议
            suggestion_para = doc.add_paragraph()
            suggestion_para.add_run('应对建议: ').bold = True
            doc.add_paragraph(risk['suggestion'])
            
            # 相关法规
            if risk['statutes']:
                statutes_para = doc.add_paragraph()
                statutes_para.add_run('相关法规: ').bold = True
                for statute in risk['statutes']:
                    doc.add_paragraph(f"• {statute['ref']}", style='List Bullet')
            
            # 添加分隔线
            if i < len(review_data['risks']):
                doc.add_paragraph('_' * 50)
        
        # 保存文档
        filename = f"contract_review_{review_data['task_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.export_dir, filename)
        doc.save(filepath)
        
        logger.info(f"DOCX report generated: {filepath}")
        return filepath
    
    def _convert_to_pdf(self, docx_path: str) -> str:
        """将DOCX转换为PDF（需要额外的库）"""
        # 这里需要使用如python-docx2pdf或其他库
        # 简化实现，实际项目中需要安装相应依赖
        try:
            from docx2pdf import convert
            pdf_path = docx_path.replace('.docx', '.pdf')
            convert(docx_path, pdf_path)
            logger.info(f"PDF report generated: {pdf_path}")
            return pdf_path
        except ImportError:
            logger.warning("docx2pdf not available, returning DOCX file")
            return docx_path
        except Exception as e:
            logger.error(f"Error converting to PDF: {e}")
            return docx_path
    
    def _get_rgb_color(self, r: int, g: int, b: int):
        """获取RGB颜色对象"""
        from docx.shared import RGBColor
        return RGBColor(r, g, b)
    
    def generate_simple_report(self, review_data: Dict[str, Any]) -> str:
        """生成简化版报告（纯文本）"""
        report_lines = []
        report_lines.append("=" * 50)
        report_lines.append("合同风险审查报告")
        report_lines.append("=" * 50)
        report_lines.append("")
        
        # 基本信息
        report_lines.append("基本信息:")
        report_lines.append(f"任务ID: {review_data['task_id']}")
        report_lines.append(f"合同类型: {review_data['contract_type']}")
        report_lines.append(f"审查角色: {review_data['role']}")
        report_lines.append(f"审查状态: {review_data['status']}")
        report_lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_lines.append("")
        
        # 风险概览
        summary = review_data['summary']
        report_lines.append("风险概览:")
        report_lines.append(f"总风险数: {summary['total_risks']}")
        report_lines.append(f"高风险: {summary['high_risks']}")
        report_lines.append(f"中风险: {summary['medium_risks']}")
        report_lines.append(f"低风险: {summary['low_risks']}")
        report_lines.append("")
        
        # 详细风险
        report_lines.append("详细风险分析:")
        report_lines.append("-" * 30)
        
        for i, risk in enumerate(review_data['risks'], 1):
            report_lines.append(f"{i}. {risk['title']}")
            report_lines.append(f"   风险级别: {risk['risk_level']}")
            if risk['clause_id']:
                report_lines.append(f"   条款编号: {risk['clause_id']}")
            report_lines.append(f"   风险描述: {risk['summary']}")
            report_lines.append(f"   应对建议: {risk['suggestion']}")
            
            if risk['statutes']:
                report_lines.append("   相关法规:")
                for statute in risk['statutes']:
                    report_lines.append(f"     • {statute['ref']}")
            
            report_lines.append("")
        
        # 保存文本报告
        filename = f"contract_review_{review_data['task_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.join(self.export_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_lines))
        
        logger.info(f"Text report generated: {filepath}")
        return filepath

# 全局导出服务实例
export_service = ExportService()